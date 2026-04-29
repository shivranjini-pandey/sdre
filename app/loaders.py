"""
Document loaders for various file types.

Handles PDF, DOCX, TXT and prepares them for chunking.
"""

from pathlib import Path
from typing import List, Tuple
import structlog # type: ignore
from pypdf import PdfReader # type: ignore
from docx import Document as DocxDocument # type: ignore

logger = structlog.get_logger()

class DocumentLoader:
    """Load documents from various formats."""
    
    @staticmethod
    def load_pdf(file_path: str) -> Tuple[str, dict]:
        """
        Load PDF file.
        
        Returns:
            (full_text, metadata)
        """
        logger.info("Loading PDF", file_path=file_path)
        
        reader = PdfReader(file_path)
        text = ""
        page_count = len(reader.pages)
        
        for page_num, page in enumerate(reader.pages):
            text += f"\n--- Page {page_num + 1} ---\n"
            text += page.extract_text()
        
        metadata = {
            "file_type": "pdf",
            "page_count": page_count,
        }
        
        return text, metadata
    
    @staticmethod
    def load_docx(file_path: str) -> Tuple[str, dict]:
        """Load DOCX file."""
        logger.info("Loading DOCX", file_path=file_path)
        
        doc = DocxDocument(file_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        
        metadata = {
            "file_type": "docx",
            "paragraph_count": len(doc.paragraphs),
        }
        
        return text, metadata
    
    @staticmethod
    def load_txt(file_path: str) -> Tuple[str, dict]:
        """Load TXT file."""
        logger.info("Loading TXT", file_path=file_path)
        
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
        
        metadata = {
            "file_type": "txt",
            "line_count": len(text.split('\n')),
        }
        
        return text, metadata
    
    @staticmethod
    def load(file_path: str) -> Tuple[str, dict]:
        """
        Auto-detect file type and load.
        
        Args:
            file_path: Path to file
            
        Returns:
            (full_text, metadata)
        """
        path = Path(file_path)
        suffix = path.suffix.lower()
        
        if suffix == '.pdf':
            return DocumentLoader.load_pdf(file_path)
        elif suffix == '.docx':
            return DocumentLoader.load_docx(file_path)
        elif suffix == '.txt':
            return DocumentLoader.load_txt(file_path)
        else:
            raise ValueError(f"Unsupported file type: {suffix}")

class DocumentChunker:
    """Split documents into chunks for retrieval."""
    
    def __init__(self, chunk_size: int = 500, overlap: int = 50):
        """
        Initialize chunker.
        
        Args:
            chunk_size: Target chunk size in tokens (approx)
            overlap: Overlap between chunks
        """
        self.chunk_size = chunk_size
        self.overlap = overlap
    
    def chunk(self, text: str) -> List[str]:
        """
        Split text into overlapping chunks.
        
        Args:
            text: Full document text
            
        Returns:
            List of text chunks
        """
        # Simple split by sentences
        sentences = text.replace('\n', ' ').split('. ')
        chunks = []
        current_chunk = ""
        
        for sentence in sentences:
            # Rough token estimate: 1 word ≈ 1.3 tokens
            tokens = len(sentence.split()) * 1.3
            
            if len(current_chunk.split()) * 1.3 + tokens > self.chunk_size:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                current_chunk = sentence
            else:
                current_chunk += sentence + ". "
        
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        # Add overlap (last n% of chunk becomes start of next)
        overlapped = []
        for i, chunk in enumerate(chunks):
            if i > 0:
                # Add last overlap% of previous chunk
                prev_chunk = chunks[i-1]
                prev_words = prev_chunk.split()
                overlap_count = max(1, len(prev_words) // (100 // self.overlap))
                overlap_text = ' '.join(prev_words[-overlap_count:])
                chunk = overlap_text + ' ' + chunk
            
            overlapped.append(chunk)
        
        logger.info(
            "Document chunked",
            original_sentences=len(sentences),
            chunks=len(overlapped),
            avg_size=sum(len(c.split()) for c in overlapped) // len(overlapped) if overlapped else 0
        )
        
        return overlapped
